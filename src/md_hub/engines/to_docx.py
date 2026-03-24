"""DOCX engine — converts MD to styled DOCX via pandoc + python-docx.

Key lessons applied:
  - NEVER copy.deepcopy Document objects -> use docxcompose
  - Emu(0) for zero indent, not None
  - Set colors on both style-level AND run-level
  - Clean fontTable embed references when using reference templates
"""

import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu, Pt, RGBColor


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_pandoc() -> str:
    """Find pandoc in PATH or known Windows location."""
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        return "pandoc"
    except (FileNotFoundError, subprocess.CalledProcessError):
        p = Path.home() / "AppData" / "Local" / "Pandoc" / "pandoc.exe"
        if p.exists():
            return str(p)
    raise FileNotFoundError("pandoc not found")


def _parse_color(hex_str: str) -> RGBColor:
    """Parse '#RRGGBB' or 'RRGGBB' to RGBColor."""
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _parse_pt(size_str: str) -> Pt:
    """Parse '11pt' or '10.5pt' to Pt."""
    return Pt(float(size_str.replace("pt", "")))


def _clean_template(docx_path: str) -> Path:
    """Strip embedded font files and references from template."""
    clean = Path(tempfile.mktemp(suffix=".docx"))
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(str(clean), "w") as zout:
            for item in zin.namelist():
                if item.startswith("word/fonts/") and item.endswith(".ttf"):
                    continue
                data = zin.read(item)
                if item == "word/_rels/fontTable.xml.rels":
                    t = data.decode("utf-8")
                    t = re.sub(r'<Relationship[^>]*Target="fonts/[^"]*\.ttf"[^/]*/>', "", t)
                    data = t.encode("utf-8")
                elif item == "[Content_Types].xml":
                    t = data.decode("utf-8")
                    t = re.sub(r'<Override[^>]*PartName="/word/fonts/[^"]*"[^/]*/>', "", t)
                    data = t.encode("utf-8")
                elif item == "word/fontTable.xml":
                    t = data.decode("utf-8")
                    t = re.sub(r'<w:embedRegular[^/]*/>', "", t)
                    t = re.sub(r'<w:embedBold[^/]*/>', "", t)
                    t = re.sub(r'<w:embedItalic[^/]*/>', "", t)
                    t = re.sub(r'<w:embedBoldItalic[^/]*/>', "", t)
                    data = t.encode("utf-8")
                zout.writestr(item, data)
    return clean


# ---------------------------------------------------------------------------
# Pandoc conversion
# ---------------------------------------------------------------------------

def _pandoc_convert(md_path: str, output_path: str, reference_docx: str = None) -> None:
    """Convert MD to DOCX via pandoc subprocess."""
    pandoc = _find_pandoc()
    cmd = [pandoc, md_path, "-o", output_path]
    if reference_docx:
        clean_ref = _clean_template(reference_docx)
        cmd.extend(["--reference-doc", str(clean_ref)])
    subprocess.run(cmd, check=True)
    if reference_docx:
        clean_ref.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Style application
# ---------------------------------------------------------------------------

def _apply_base_styles(doc: Document, config: dict) -> None:
    """Apply font and color styles from config at style-level and run-level."""
    font_cfg = config.get("font", {})
    colors_cfg = config.get("colors", {})

    heading_font = font_cfg.get("heading", "Malgun Gothic")
    body_font = font_cfg.get("body", "Malgun Gothic")
    heading_color = _parse_color(colors_cfg.get("heading", "#1A1A1A"))
    body_color = _parse_color(colors_cfg.get("body", "#333333"))

    # Style-level: Normal
    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal.font.color.rgb = body_color

    # Style-level: Headings
    HEADING_SIZES = {
        "Heading 1": Pt(14),
        "Heading 2": Pt(12),
        "Heading 3": Pt(11),
    }
    for sname, size in HEADING_SIZES.items():
        try:
            style = doc.styles[sname]
            style.font.name = heading_font
            style.font.size = size
            style.font.bold = True if sname != "Heading 3" else False
            style.font.color.rgb = heading_color
            style.paragraph_format.line_spacing = 1.15
            # Emu(0) not None — None removes XML element, Word falls back
            style.paragraph_format.left_indent = Emu(0)
            style.paragraph_format.first_line_indent = Emu(0)
        except KeyError:
            pass

    # Run-level: set colors on every run (run overrides style)
    for para in doc.paragraphs:
        sname = para.style.name
        if sname in HEADING_SIZES:
            for run in para.runs:
                run.font.color.rgb = heading_color
                run.font.name = heading_font
        else:
            for run in para.runs:
                run.font.color.rgb = body_color
                run.font.name = body_font


def _apply_table_styles(doc: Document, config: dict) -> None:
    """Apply table styling: header bg, borders, 100% width."""
    colors_cfg = config.get("colors", {})
    header_bg = colors_cfg.get("table_header_bg", "F2F2F2").lstrip("#")

    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Width 100%
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        # Borders: full grid
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        ))

        # Per-cell styling
        for i, row in enumerate(table.rows):
            is_header = (i == 0)
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                old_shd = tcPr.find(qn("w:shd"))
                if old_shd is not None:
                    tcPr.remove(old_shd)
                if is_header:
                    tcPr.append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{header_bg}" w:val="clear"/>'
                    ))
                for p in cell.paragraphs:
                    if is_header:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        if is_header:
                            run.font.bold = True


# ---------------------------------------------------------------------------
# Heading numbering
# ---------------------------------------------------------------------------

def _apply_heading_numbering(doc: Document, annex_numbering: bool = True) -> None:
    """Inject numPr XML for heading levels. Skip Annex sections if annex_numbering=False."""
    ILVL = {"Heading 1": "0", "Heading 2": "1", "Heading 3": "2"}
    in_annex = False

    for para in doc.paragraphs:
        if para.style.name == "Heading 1":
            if para.text.strip().startswith("Annex"):
                in_annex = True
            else:
                in_annex = False

        if in_annex and not annex_numbering:
            continue

        if para.style.name in ILVL:
            pPr = para._p.get_or_add_pPr()
            old = pPr.find(qn("w:numPr"))
            if old is not None:
                pPr.remove(old)
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), ILVL[para.style.name])
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "1")
            numPr.append(ilvl)
            numPr.append(numId)
            # Insert after pStyle (schema order)
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                pStyle.addnext(numPr)
            else:
                pPr.insert(0, numPr)


# ---------------------------------------------------------------------------
# Header/Footer update
# ---------------------------------------------------------------------------

def _update_doc_code(doc: Document, doc_code: str) -> None:
    """Update header/footer with doc_code if present."""
    for section in doc.sections:
        if section.header and section.header.paragraphs:
            for p in section.header.paragraphs:
                if p.text:
                    p.clear()
                    run = p.add_run(doc_code)
                    run.font.size = Pt(9)
        if section.footer and section.footer.paragraphs:
            for p in section.footer.paragraphs:
                if p.text:
                    p.clear()
                    run = p.add_run(doc_code)
                    run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def convert(md_path: str, output_path: str, template_name: str = "default", **kwargs) -> None:
    """Convert MD to styled DOCX.

    Args:
        md_path: Path to input Markdown file.
        output_path: Path for output DOCX file.
        template_name: Template name to load from templates.yaml.
        **kwargs: Optional overrides (doc_code, etc.)
    """
    from md_hub.template_loader import load_template
    config = load_template(template_name)

    tmp_docx = tempfile.mktemp(suffix=".docx")

    # Step 1-2: pandoc MD -> DOCX
    _pandoc_convert(md_path, tmp_docx)

    # Step 3: If cover template, merge via docxcompose
    cover = config.get("cover")
    if cover:
        from docxcompose.composer import Composer
        clean = _clean_template(str(cover))
        master = Document(str(clean))
        clean.unlink(missing_ok=True)
        composer = Composer(master)
        composer.append(Document(tmp_docx))
        doc = master
        Path(tmp_docx).unlink(missing_ok=True)
    else:
        doc = Document(tmp_docx)
        Path(tmp_docx).unlink(missing_ok=True)

    # Step 4: Apply base styles
    _apply_base_styles(doc, config)

    # Step 5: Apply table styles
    _apply_table_styles(doc, config)

    # Step 6: Heading numbering (conditional)
    if config.get("heading_numbering"):
        annex_num = config.get("annex_numbering", True)
        _apply_heading_numbering(doc, annex_numbering=annex_num)

    # Step 7: Doc code (conditional)
    doc_code = kwargs.get("doc_code")
    if doc_code:
        _update_doc_code(doc, doc_code)

    doc.save(output_path)
