"""Two-stage MD→docx converter for KHC SOP documents.

Stage 1: Preprocess Google Docs MD → standard MD, then pandoc → body docx
Stage 2: docxcompose to merge body into template (no copy.deepcopy)
         + numPr injection for heading numbering

Key insight: copy.deepcopy between Document objects causes corruption.
docxcompose handles style/numbering conflicts properly.
"""

import argparse
import re
import subprocess
import tempfile
from pathlib import Path
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docxcompose.composer import Composer


# ---------------------------------------------------------------------------
# Stage 1: Preprocess + Pandoc
# ---------------------------------------------------------------------------

def preprocess_md(text: str) -> tuple[dict, str]:
    """Split MD into metadata (cover/approval/revision) and body content."""
    lines = text.split("\n")
    metadata = {
        "title": "", "company": "",
        "approval": [], "revision": [],
        "version": "", "effective_date": "",
    }
    body_lines = []
    state = "cover"
    table_buf = []

    for line in lines:
        stripped = line.strip()

        if state == "cover":
            if re.match(r"^\|.*\|$", stripped) and ":---:" not in stripped:
                t = stripped.strip("|").strip()
                if t and not t.startswith(":"):
                    metadata["title"] = t
            if "카카오헬스케어" in stripped:
                metadata["company"] = stripped.strip("*").strip()
            if "작성 · 검토 · 승인 이력" in stripped:
                state = "approval"
                table_buf = []
            continue

        if state == "approval":
            if stripped.startswith("|"):
                table_buf.append(stripped)
            elif "제" in stripped and "개정" in stripped:
                metadata["approval"] = _parse_table(table_buf)
                state = "revision"
                table_buf = []
            continue

        if state == "revision":
            if stripped.startswith("|"):
                table_buf.append(stripped)
            elif "목차" in stripped:
                metadata["revision"] = _parse_table(table_buf)
                if metadata["revision"]:
                    last = metadata["revision"][-1]
                    if len(last) >= 3:
                        metadata["version"] = last[0]
                        metadata["effective_date"] = last[2] or last[1]
                state = "toc"
            continue

        if state == "toc":
            if re.match(r"^\s*(?:\d+\.?\s+)?#{1,4}\s+", stripped):
                state = "body"
            else:
                continue

        if state == "body":
            # Clean Google Docs heading: "1. # **Text** {#id}" → "# Text"
            hm = re.match(
                r"^\s*(?:\d+\.?\s+)?(#{1,4})\s+\*{0,2}(.*?)\*{0,2}\s*(?:\{#[^}]*\})?\s*$",
                stripped
            )
            if hm:
                hashes = hm.group(1)
                htxt = hm.group(2).strip().strip("*").replace("**", "")
                body_lines.append(f"{hashes} {htxt}")
                continue
            body_lines.append(line)

    return metadata, "\n".join(body_lines)


def _parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        s = line.strip()
        if not s.startswith("|"):
            continue
        inner = s.strip("|").strip()
        if all(c in "-: |" for c in inner) and "-" in inner:
            continue
        rows.append([c.strip() for c in s.strip("|").split("|")])
    return rows


def pandoc_body(body_md: str, reference_docx: str) -> Path:
    """Pandoc converts body MD to docx using clean reference template."""
    tmp_md = Path(tempfile.mktemp(suffix=".md"))
    tmp_docx = Path(tempfile.mktemp(suffix=".docx"))
    tmp_md.write_text(body_md, encoding="utf-8")

    pandoc = _find_pandoc()
    clean_ref = _clean_template(reference_docx)

    subprocess.run([pandoc, str(tmp_md), "--reference-doc", str(clean_ref),
                    "-o", str(tmp_docx)], check=True)

    tmp_md.unlink()
    clean_ref.unlink()
    return tmp_docx


def _find_pandoc() -> str:
    try:
        subprocess.run(["pandoc", "--version"], capture_output=True, check=True)
        return "pandoc"
    except (FileNotFoundError, subprocess.CalledProcessError):
        p = Path.home() / "AppData" / "Local" / "Pandoc" / "pandoc.exe"
        if p.exists():
            return str(p)
    raise FileNotFoundError("pandoc not found")


def _clean_template(docx_path: str) -> Path:
    """Strip embedded font files AND all references from template."""
    clean = Path(tempfile.mktemp(suffix=".docx"))
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(str(clean), "w") as zout:
            for item in zin.namelist():
                # Skip actual font files
                if item.startswith("word/fonts/") and item.endswith(".ttf"):
                    continue
                data = zin.read(item)
                if item == "word/_rels/fontTable.xml.rels":
                    t = data.decode("utf-8")
                    t = re.sub(r'<Relationship[^>]*Target="fonts/[^"]*\.ttf"[^/]*/>', "", t)
                    data = t.encode("utf-8")
                elif item == "[Content_Types].xml":
                    t = data.decode("utf-8")
                    t = re.sub(r'<Override[^>]*PartName="/word/fonts/[^"]*"[^/]*/>', "", t)
                    data = t.encode("utf-8")
                elif item == "word/fontTable.xml":
                    # Remove embedRegular/embedBold elements from font definitions
                    t = data.decode("utf-8")
                    t = re.sub(r'<w:embedRegular[^/]*/>', "", t)
                    t = re.sub(r'<w:embedBold[^/]*/>', "", t)
                    t = re.sub(r'<w:embedItalic[^/]*/>', "", t)
                    t = re.sub(r'<w:embedBoldItalic[^/]*/>', "", t)
                    data = t.encode("utf-8")
                zout.writestr(item, data)
    return clean


# ---------------------------------------------------------------------------
# Stage 2: Template + docxcompose merge
# ---------------------------------------------------------------------------

def build_cover_doc(metadata: dict, reference_docx: str) -> Path:
    """Create a small docx with just the cover page elements."""
    clean = _clean_template(reference_docx)
    doc = Document(str(clean))
    clean.unlink()

    # Clear body
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    # Title table
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata.get("title", "Untitled"))
    run.bold = True
    run.font.size = Pt(22)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    ))

    # Company
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(metadata.get("company", "주식회사 카카오헬스케어 데이터부문"))
    r.bold = True
    r.font.size = Pt(16)

    # Approval + Revision tables
    for label, rows in [("작성 · 검토 · 승인 이력", metadata.get("approval", [])),
                         ("제 · 개정 이력", metadata.get("revision", []))]:
        lp = doc.add_paragraph()
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(12)
        if rows:
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, ct in enumerate(row):
                    if j < ncols:
                        c = tbl.rows[i].cells[j]
                        tp = c.paragraphs[0]
                        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        tr = tp.add_run(ct)
                        tr.font.size = Pt(10)
                        if i == 0:
                            tr.bold = True

    # Page break
    bp = doc.add_paragraph()
    br_run = bp.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    br_run._r.append(br)

    cover_path = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(str(cover_path))
    return cover_path


def apply_heading_numbering(doc: Document):
    """Add numPr to heading paragraphs: numId=1, ilvl=heading_level-1."""
    ILVL = {"Heading 1": "0", "Heading 2": "1", "Heading 3": "2"}
    in_annex = False
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and para.text.strip().startswith("Annex"):
            in_annex = True
        elif para.style.name == "Heading 1" and not para.text.strip().startswith("Annex"):
            in_annex = False
        if in_annex:
            continue
        if para.style.name in ILVL:
            pPr = para._p.get_or_add_pPr()
            # Remove existing numPr
            old = pPr.find(qn("w:numPr"))
            if old is not None:
                pPr.remove(old)
            # Add numPr with correct element order
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), ILVL[para.style.name])
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "1")
            numPr.append(ilvl)
            numPr.append(numId)
            # Insert after pStyle (schema order: pStyle → numPr → ...)
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                pStyle.addnext(numPr)
            else:
                pPr.insert(0, numPr)


def fix_numbering_indents(doc: Document):
    """Rewrite abstractNum level indents to match our desired hierarchy.

    Target (twips, 1cm = 567 twips):
      lvl 0 (H1): left=284 (0.5cm), hanging=284 → text starts at 0cm+number
      lvl 1 (H2): left=567 (1.0cm), hanging=284 → text starts at 0.5cm+number
      lvl 2 (H3): left=851 (1.5cm), hanging=284 → text starts at 1.0cm+number

    Bullet numbering (abstractNum 2):
      lvl 0: left=567 (1.0cm), hanging=284
      lvl 1: left=851 (1.5cm), hanging=284
      lvl 2: left=1134 (2.0cm), hanging=284
    """
    import zipfile as _zf
    numbering_part = doc.part.numbering_part
    if numbering_part is None:
        return

    numbering_elem = numbering_part._element

    NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Heading numbering indents (twips: 567 = 1cm)
    # left = where heading text starts, hanging = number width
    # number starts at (left - hanging)
    HEADING_INDENTS = {
        "0": {"left": "0", "hanging": "0"},       # H1: starts at margin
        "1": {"left": "227", "hanging": "0"},     # H2: starts at 0.4cm
        "2": {"left": "454", "hanging": "0"},     # H3: starts at 0.8cm
    }

    for an in numbering_elem.findall(qn("w:abstractNum")):
        anid = an.get(qn("w:abstractNumId"))
        if anid == "1":  # Heading numbering
            for lvl in an.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                if ilvl in HEADING_INDENTS:
                    pPr = lvl.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        lvl.insert(0, pPr)
                    ind = pPr.find(qn("w:ind"))
                    if ind is None:
                        ind = OxmlElement("w:ind")
                        pPr.append(ind)
                    vals = HEADING_INDENTS[ilvl]
                    ind.set(qn("w:left"), vals["left"])
                    ind.set(qn("w:hanging"), vals["hanging"])


def fix_formatting(doc: Document):
    """Apply formatting at PARAGRAPH and RUN level (overrides styles).

    Formatting:
      H1: 0cm indent, 13pt bold, line 1.15, before 12pt / after 4pt
      H2: 0.3cm indent, 11pt bold, line 1.15, before 8pt / after 3pt
      H3: 0.6cm indent, 10.5pt, line 1.15, before 6pt / after 2pt
      Body after H1: 0.3cm indent, 10pt, line 1.15
      Body after H2: 0.6cm indent, 10pt, line 1.15
      Body after H3: 0.9cm indent, 10pt, line 1.15
      Bullet: body indent + 0.3cm, 10pt
    """
    from docx.shared import Cm

    # body_indent must match numbering "left" value exactly
    # 284 twips = 0.501cm, 567 twips = 1.0cm, 851 twips = 1.501cm
    from docx.shared import Emu
    TWIP = 635  # 1 twip = 635 EMU

    # B-style: all body text aligns to H1 baseline (left margin)
    HEADING_FMT = {
        "Heading 1": {"size": Pt(14), "bold": True,
                      "sb": Pt(12), "sa": Pt(6), "body_indent": Emu(0)},
        "Heading 2": {"size": Pt(12), "bold": True,
                      "sb": Pt(9), "sa": Pt(3), "body_indent": Emu(0)},
        "Heading 3": {"size": Pt(11), "bold": False,
                      "sb": Pt(6), "sa": Pt(3), "body_indent": Emu(0)},
    }
    BULLET_EXTRA = Cm(0.5)
    BODY_SIZE = Pt(10)

    # Set style-level: font size (for numbering), force indent to 0 (not None!)
    # None = remove element = Word falls back to built-in "제목 3" default
    # 0 = explicit override = no style-level indent, numbering controls it
    from docx.shared import Emu
    for sname, fmt in HEADING_FMT.items():
        try:
            style = doc.styles[sname]
            style.font.size = fmt["size"]
            style.font.bold = fmt["bold"]
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.left_indent = Emu(0)
            style.paragraph_format.first_line_indent = Emu(0)
            style.paragraph_format.space_before = fmt["sb"]
            style.paragraph_format.space_after = fmt["sa"]
        except KeyError:
            pass

    current_body_indent = None
    in_body = False

    for para in doc.paragraphs:
        sname = para.style.name

        # --- Annex headings: no numbering, page break for B/C onwards ---
        is_annex_h1 = (sname == "Heading 1"
                       and para.text.strip().startswith("Annex"))
        if is_annex_h1:
            in_annex = True
            # Page break before Annex B, C, ... (not A)
            if not para.text.strip().startswith("Annex A"):
                pf = para.paragraph_format
                pf.page_break_before = True

        if not hasattr(fix_formatting, '_annex') and is_annex_h1:
            in_annex = True
        if is_annex_h1:
            in_annex = True
        elif sname == "Heading 1":
            in_annex = False

        # --- Headings: numbering controls indent, clear paragraph/style indent ---
        if sname in HEADING_FMT:
            fmt = HEADING_FMT[sname]
            pf = para.paragraph_format
            # Force 0 (not None!) so Word doesn't fall back to built-in style default
            pf.left_indent = Emu(0)
            pf.first_line_indent = Emu(0)
            pf.space_before = fmt["sb"]
            pf.space_after = fmt["sa"]
            pf.line_spacing = 1.15
            for run in para.runs:
                run.font.size = fmt["size"]
                run.font.bold = fmt["bold"]
                run.font.color.rgb = RGBColor(0, 0, 0)
            current_body_indent = fmt["body_indent"]
            in_body = True
            continue

        # --- Body text: direct formatting ---
        if not in_body or current_body_indent is None:
            continue

        # Detect bullet
        is_bullet = sname in ("List Paragraph", "List Bullet",
                               "Compact", "List Number")
        if not is_bullet:
            pPr = para._p.find(qn("w:pPr"))
            if pPr is not None:
                np = pPr.find(qn("w:numPr"))
                if np is not None:
                    nid = np.find(qn("w:numId"))
                    if nid is not None and nid.get(qn("w:val")) != "1":
                        is_bullet = True

        pf = para.paragraph_format
        pf.line_spacing = 1.15
        pf.space_after = Pt(2)

        if is_bullet:
            pf.left_indent = current_body_indent + BULLET_EXTRA
        else:
            pf.left_indent = current_body_indent

        for run in para.runs:
            run.font.size = BODY_SIZE


def apply_table_style(doc: Document):
    """Apply standard KHC table style to all tables.

    Style spec (A01 Table 0 baseline):
      - Header row (row 0): #F2F2F2 background, bold, center aligned
      - Data rows: no background, first column center aligned
      - Borders: single, sz=4, full grid
      - Width: 100% (pct 5000)
    """
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Width 100%
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        # Borders: full grid
        old_borders = tblPr.find(qn("w:tblBorders"))
        if old_borders is not None:
            tblPr.remove(old_borders)
        tblPr.append(parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        ))

        # Per-cell styling
        for i, row in enumerate(table.rows):
            is_header = (i == 0)
            for j, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()

                # Shading
                old_shd = tcPr.find(qn("w:shd"))
                if old_shd is not None:
                    tcPr.remove(old_shd)
                if is_header:
                    tcPr.append(parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'
                    ))

                # Alignment & font
                for p in cell.paragraphs:
                    if is_header:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif j == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        if is_header:
                            run.font.bold = True


def update_header(doc: Document, title: str, doc_code: str, version: str):
    for section in doc.sections:
        if section.header.tables:
            cell = section.header.tables[0].rows[0].cells[-1]
            for p in cell.paragraphs:
                p.clear()
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{title}\n{doc_code} Ver {version}")
            run.font.size = Pt(9)


def update_footer(doc: Document, effective_date: str):
    for section in doc.sections:
        if section.footer.tables:
            cell = section.footer.tables[0].rows[0].cells[0]
            for p in cell.paragraphs:
                p.clear()
            p = cell.paragraphs[0]
            run = p.add_run(f"Effective date : {effective_date}")
            run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def convert(template_path: str, input_path: str, output_path: str, doc_code: str = ""):
    md_text = Path(input_path).read_text(encoding="utf-8")
    metadata, body_md = preprocess_md(md_text)

    if not doc_code:
        m = re.match(r"(DP\d+(?:-A\d+)?)", Path(input_path).stem)
        doc_code = m.group(1) if m else ""

    version = metadata.get("version", "1.0").replace("V", "").replace("v", "")

    # Stage 1: pandoc body
    print(f"Stage 1: pandoc body ({len(body_md)} chars)...")
    body_docx = pandoc_body(body_md, template_path)

    # Stage 2: build cover doc
    print("Stage 2: building cover...")
    cover_docx = build_cover_doc(metadata, template_path)

    # Stage 3: merge with docxcompose (cover = master, body = appended)
    print("Stage 3: docxcompose merge...")
    master = Document(str(cover_docx))
    composer = Composer(master)
    composer.append(Document(str(body_docx)))

    # Apply heading numbering + fix numbering indent
    print("Stage 4: applying heading numbering...")
    apply_heading_numbering(master)
    fix_numbering_indents(master)

    # Fix line spacing and indentation
    print("Stage 5: fixing line spacing + indentation...")
    fix_formatting(master)

    # Apply table style
    print("Stage 6: applying table style...")
    apply_table_style(master)

    # Update header/footer
    update_header(master, metadata.get("title", ""), doc_code, version)
    update_footer(master, metadata.get("effective_date", ""))

    master.save(output_path)
    print(f"Saved: {output_path}")

    # Cleanup
    cover_docx.unlink()
    body_docx.unlink()


def main():
    parser = argparse.ArgumentParser(description="KHC SOP MD→docx (docxcompose)")
    parser.add_argument("--template", required=True)
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--doc-code", default="")
    args = parser.parse_args()
    convert(args.template, args.input, args.output, args.doc_code)


if __name__ == "__main__":
    main()
